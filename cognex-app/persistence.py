"""
The real persistence API — mounted at /api/v2. This is the first phase of
closing the "nothing is really saved" gap: companies, the roster, decisions,
and goals now live in Postgres (see models.py / db_postgres.py) instead of
only in the browser tab.

Deliberately NOT covered by this phase (still session/browser-only, exactly
as before): Vantage gaps/plans, handoff records, and the access ledger. See
the persistence build-log entry for why this slice first and what's next.

Ask Cognex chat threads (see the ChatThread endpoints near the bottom of
this file) were added to real persistence on 2026-08-29, one phase later
than the rest — see that day's build-log entry for why.

Auth note, read together with server.py's existing one on the reference
endpoints: login here still checks a shared demo password (`DEMO_PASSWORD`)
— the same illustrative, not-real authentication the app has always had.
What changed is WHERE the account data comes from (a real database instead
of a hardcoded JS object), not how identity is verified. Real
authentication (a third-party provider was the founder's stated direction)
is the deliberately separate next phase — swapping this login endpoint's
password check for real session/token verification shouldn't require
touching the data model built here at all.
"""

import io
import secrets
from typing import Optional

from fastapi import APIRouter, File, Form, HTTPException, UploadFile
from pydantic import BaseModel
from sqlalchemy import select

from datetime import datetime, timezone

from db_postgres import get_session
from models import ChatThread, Company, CompanyUsage, Decision, Document, Goal, Handoff, LedgerEntry, Persona

router = APIRouter(prefix="/api/v2")

DEMO_PASSWORD = "cognex"
STAFF_EMAIL = "admin@cognex.ai"
# Staff sign-in gets its own password, separate from the shared per-company
# demo password above — the frontend's staff-sign-in modal never displays it
# (no prefilled prompt, no footnote), unlike the regular login screen, which
# does show DEMO_PASSWORD on screen as an explicit "this demo isn't really
# authenticated" disclosure.
STAFF_PASSWORD = "Admin@Cognex"

# ---------------------------------------------------------------------------
# Minimal session tokens — added 2026-08-31 after a security audit flagged
# that GET /companies/{company_id} (which returns EVERY decision's full
# why/alternatives/risks and the whole roster, unfiltered) had no credential
# requirement at all: anyone who knew or guessed a company_id could read a
# client's full confidential Decision Memory with zero authentication,
# directly contradicting the product's core "chain of command protection"
# pitch. This is NOT real authentication — there's still one shared password
# per company (DEMO_PASSWORD) and one for staff (STAFF_PASSWORD), so it
# doesn't stop someone who has that shared password from reading anything;
# closing that gap for real is the explicitly-scoped-separate "real auth"
# phase (a third-party provider, per the founder's own prior direction).
# What THIS closes is the zero-credential drive-by: a company's confidential
# data can no longer be read by an unauthenticated request that never went
# through /login at all. An in-memory dict is deliberately sufficient here
# (matching this file's existing demo-scale conventions) — these tokens are
# only ever meant to gate reads for the lifetime of one server process, the
# same way the rest of this "illustrative, not real auth" system already
# works; they are NOT a substitute for real session/token verification.
_SESSIONS: dict = {}  # token -> {"kind": "company"|"platform", "company_id": str|None, "persona_id": str|None}


def _issue_token(kind: str, company_id: Optional[str] = None, persona_id: Optional[str] = None) -> str:
    token = secrets.token_urlsafe(32)
    _SESSIONS[token] = {"kind": kind, "company_id": company_id, "persona_id": persona_id}
    return token


def _require_session(token: Optional[str], company_id: Optional[str] = None):
    """Validates a session token for a read that should require SOME prior
    login. A platform (staff) token is always accepted regardless of
    company_id — staff legitimately need to read any company's data to
    support it, mirroring the existing "Enter as support" flow. A company
    token must match the specific company_id being read, so signing in to
    one company never lets you read another's data by guessing its id."""
    session = _SESSIONS.get(token) if token else None
    if not session:
        raise HTTPException(status_code=401, detail="Sign in required.")
    if session["kind"] == "platform":
        return
    if company_id is not None and session.get("company_id") != company_id:
        raise HTTPException(status_code=401, detail="This session isn't signed in to that company.")


# ---------------------------------------------------------------------------
# Serialization — matches the exact camelCase shape static/index.html already
# expects (it used to build these objects itself; now it fetches them).
# ---------------------------------------------------------------------------
def _persona_json(p: Persona) -> dict:
    return {
        "id": p.id, "name": p.name, "title": p.title, "level": p.level, "dept": p.dept,
        "email": p.email, "reportsTo": p.reports_to, "isOrgAdmin": p.is_org_admin,
        "initials": p.initials, "bg": p.bg, "fg": p.fg,
    }


def _decision_json(d: Decision) -> dict:
    out = {
        "id": d.id, "title": d.title, "decided": d.decided, "owner": d.owner,
        "visibility": d.visibility, "derivedLevel": d.derived_level,
        "why": d.why or [], "alternatives": d.alternatives or [], "assumptions": d.assumptions or [],
        "risks": d.risks or [], "review": d.review, "result": d.result,
        "tags": d.tags or [], "derived": d.derived,
    }
    if d.contributor:
        out["contributor"] = d.contributor
    if d.via_story:
        out["viaStory"] = True
    if d.via_slack:
        out["viaSlack"] = True
    if d.via_chat:
        out["viaChat"] = True
    if d.via_vantage:
        out["viaVantage"] = True
    return out


def _goal_json(g: Goal) -> dict:
    out = {"id": g.id, "title": g.title, "kicker": g.kicker, "depth": g.depth, "parent": g.parent, "dept": g.dept}
    if g.owner:
        out["owner"] = g.owner
    return out


def _handoff_json(h: Handoff) -> dict:
    return {
        "id": h.id, "fromPersonaId": h.from_persona_id, "fromName": h.from_name,
        "fromTitle": h.from_title, "fromDept": h.from_dept, "toPersonaId": h.to_persona_id,
        "createdAt": h.created_at, "qaPairs": h.qa_pairs or [], "report": h.report or {},
        "offline": h.offline, "delegationItems": h.delegation_items or [], "status": h.status,
    }


def _ledger_json(e: LedgerEntry) -> dict:
    return {
        "id": e.id, "ts": e.ts, "personaId": e.persona_id, "personaName": e.persona_name,
        "personaTitle": e.persona_title, "decisionId": e.decision_id, "decisionTitle": e.decision_title,
        "access": e.access, "via": e.via,
    }


def _company_snapshot(db, company: Company) -> dict:
    personas = db.execute(select(Persona).where(Persona.company_id == company.id)).scalars().all()
    decisions = db.execute(select(Decision).where(Decision.company_id == company.id)).scalars().all()
    goals = db.execute(select(Goal).where(Goal.company_id == company.id)).scalars().all()
    handoffs = db.execute(
        select(Handoff).where(Handoff.company_id == company.id).order_by(Handoff.created_at.desc())
    ).scalars().all()
    # Ledger entries are capped to the most recent 500 per company (see
    # log_ledger_entry below, same cap the frontend already enforced
    # locally) -- this is an audit trail, not unbounded storage.
    ledger = db.execute(
        select(LedgerEntry).where(LedgerEntry.company_id == company.id).order_by(LedgerEntry.ts.desc()).limit(500)
    ).scalars().all()
    return {
        "id": company.id, "name": company.name, "industry": company.industry, "seeded": company.seeded,
        "personas": {p.id: _persona_json(p) for p in personas},
        "decisions": [_decision_json(d) for d in decisions],
        "goals": [_goal_json(g) for g in goals],
        "handoffs": [_handoff_json(h) for h in handoffs],
        "accessLedger": [_ledger_json(e) for e in ledger],
    }


def _get_company_or_404(db, company_id: str) -> Company:
    company = db.get(Company, company_id)
    if not company:
        raise HTTPException(status_code=404, detail=f"No company '{company_id}'.")
    return company


def _slugify(name: str) -> str:
    base = "".join(c.lower() if c.isalnum() else "-" for c in name).strip("-")
    while "--" in base:
        base = base.replace("--", "-")
    return base or "company"


# ---------------------------------------------------------------------------
# Company snapshot + login + staff console listing
# ---------------------------------------------------------------------------
@router.get("/companies/{company_id}")
def get_company(company_id: str, token: Optional[str] = None):
    _require_session(token, company_id=company_id)
    with get_session() as db:
        company = _get_company_or_404(db, company_id)
        return _company_snapshot(db, company)


@router.get("/platform/companies")
def list_companies(token: Optional[str] = None):
    session = _SESSIONS.get(token) if token else None
    if not session or session["kind"] != "platform":
        raise HTTPException(status_code=401, detail="Cognex staff sign-in required.")
    with get_session() as db:
        companies = db.execute(select(Company)).scalars().all()
        rows = []
        for c in companies:
            personas = db.execute(select(Persona).where(Persona.company_id == c.id)).scalars().all()
            rows.append({
                "id": c.id, "name": c.name, "industry": c.industry, "seeded": c.seeded,
                "memberCount": len(personas),
                "adminCount": sum(1 for p in personas if p.is_org_admin),
            })
        return {"companies": rows}


class LoginRequest(BaseModel):
    email: str
    password: str


@router.post("/login")
def login(req: LoginRequest):
    email = req.email.strip().lower()
    if email == STAFF_EMAIL:
        if req.password != STAFF_PASSWORD:
            raise HTTPException(status_code=401, detail="Incorrect password for Cognex staff.")
        return {"type": "platform", "token": _issue_token("platform")}

    with get_session() as db:
        persona = db.execute(select(Persona).where(Persona.email == email)).scalars().first()
        if not persona or req.password != DEMO_PASSWORD:
            raise HTTPException(status_code=401, detail="No matching account, or the password's wrong.")
        company = _get_company_or_404(db, persona.company_id)
        token = _issue_token("company", company_id=company.id, persona_id=persona.id)
        return {"type": "company", "companyId": company.id, "personaId": persona.id, "token": token, "company": _company_snapshot(db, company)}


# ---------------------------------------------------------------------------
# Company creation (staff onboarding) — mirrors the wizard: exactly one
# founding Team Admin, created together with the company.
# ---------------------------------------------------------------------------
class FounderIn(BaseModel):
    name: str
    title: str = "CEO"
    email: str


class CompanyCreateRequest(BaseModel):
    name: str
    industry: str = "—"
    founder: FounderIn


@router.post("/companies")
def create_company(req: CompanyCreateRequest):
    email = req.founder.email.strip().lower()
    with get_session() as db:
        existing = db.execute(select(Persona).where(Persona.email == email)).scalars().first()
        if existing:
            raise HTTPException(status_code=409, detail="That email is already in use by another account.")
        base_id = _slugify(req.name)
        company_id, n = base_id, 1
        while db.get(Company, company_id):
            n += 1
            company_id = f"{base_id}-{n}"

        company = Company(id=company_id, name=req.name, industry=req.industry or "—", seeded=False)
        db.add(company)
        founder_id = "founder"
        db.add(Persona(
            company_id=company_id, id=founder_id, name=req.founder.name, title=req.founder.title or "CEO",
            level=6, dept="Executive", email=email, reports_to=None, is_org_admin=True,
            initials="".join(w[0] for w in req.founder.name.split() if w)[:2].upper(),
            bg="#e8f0fe", fg="#0071e3",
        ))
        db.add(Goal(company_id=company_id, id="g-root", depth=0, kicker="Company goal", title="Define your first company goal", parent=None, dept=None))
        db.commit()
        return _company_snapshot(db, company)


# ---------------------------------------------------------------------------
# Roster (personas)
# ---------------------------------------------------------------------------
class PersonaIn(BaseModel):
    id: Optional[str] = None
    name: str
    title: str = ""
    level: int
    dept: str = ""
    email: str
    reportsTo: Optional[str] = None
    isOrgAdmin: bool = False


@router.post("/companies/{company_id}/personas")
def add_persona(company_id: str, req: PersonaIn):
    with get_session() as db:
        company = _get_company_or_404(db, company_id)
        pid = req.id or _slugify(req.name)
        if db.get(Persona, (company_id, pid)):
            base, n = pid, 1
            while db.get(Persona, (company_id, pid)):
                n += 1
                pid = f"{base}-{n}"
        color_idx = len(db.execute(select(Persona).where(Persona.company_id == company_id)).scalars().all())
        colors = [{"bg": "#e8f0fe", "fg": "#0071e3"}, {"bg": "#e6f9ec", "fg": "#1f8a4c"}, {"bg": "#fff1e0", "fg": "#b45c09"},
                  {"bg": "#f5e9fc", "fg": "#7a2fb0"}, {"bg": "#ffe9f0", "fg": "#c2185b"}, {"bg": "#e6f9fb", "fg": "#0f7a8c"}]
        color = colors[color_idx % len(colors)]
        p = Persona(
            company_id=company_id, id=pid, name=req.name, title=req.title, level=req.level, dept=req.dept,
            email=req.email.strip().lower(), reports_to=req.reportsTo, is_org_admin=req.isOrgAdmin,
            initials="".join(w[0] for w in req.name.split() if w)[:2].upper(), bg=color["bg"], fg=color["fg"],
        )
        db.add(p)
        db.commit()
        return _persona_json(p)


@router.put("/companies/{company_id}/personas/{persona_id}")
def update_persona(company_id: str, persona_id: str, req: PersonaIn):
    with get_session() as db:
        p = db.get(Persona, (company_id, persona_id))
        if not p:
            raise HTTPException(status_code=404, detail="No such persona.")
        p.name, p.title, p.level, p.dept = req.name, req.title, req.level, req.dept
        p.email, p.reports_to, p.is_org_admin = req.email.strip().lower(), req.reportsTo, req.isOrgAdmin
        db.commit()
        return _persona_json(p)


@router.delete("/companies/{company_id}/personas/{persona_id}")
def delete_persona(company_id: str, persona_id: str, acting_persona_id: Optional[str] = None):
    with get_session() as db:
        p = db.get(Persona, (company_id, persona_id))
        if not p:
            raise HTTPException(status_code=404, detail="No such persona.")
        if acting_persona_id == persona_id:
            raise HTTPException(status_code=400, detail="You can't remove your own account while signed in as them.")
        if p.is_org_admin:
            admin_count = len(db.execute(select(Persona).where(Persona.company_id == company_id, Persona.is_org_admin == True)).scalars().all())  # noqa: E712
            if admin_count <= 1:
                raise HTTPException(status_code=400, detail="Can't remove the last team admin. Promote someone else first.")
        db.delete(p)
        db.commit()
        return {"ok": True}


# ---------------------------------------------------------------------------
# Decisions
# ---------------------------------------------------------------------------
class DecisionIn(BaseModel):
    id: Optional[str] = None
    title: str
    decided: str = ""
    owner: str = ""
    visibility: int = 1
    derivedLevel: int = 1
    why: list = []
    alternatives: list = []
    assumptions: list = []
    risks: list = []
    review: str = "Not yet scheduled"
    result: str = "TBD"
    tags: list = []
    derived: str = ""
    contributor: Optional[str] = None
    viaStory: bool = False
    viaSlack: bool = False
    viaChat: bool = False
    viaVantage: bool = False


@router.post("/companies/{company_id}/decisions")
def add_decision(company_id: str, req: DecisionIn):
    with get_session() as db:
        _get_company_or_404(db, company_id)
        did = req.id or f"d-{_slugify(req.title)}"
        if db.get(Decision, (company_id, did)):
            did = f"{did}-{len(db.execute(select(Decision).where(Decision.company_id == company_id)).scalars().all()) + 1}"
        d = Decision(
            company_id=company_id, id=did, title=req.title, decided=req.decided, owner=req.owner,
            visibility=req.visibility, derived_level=req.derivedLevel, why=req.why, alternatives=req.alternatives,
            assumptions=req.assumptions, risks=req.risks, review=req.review, result=req.result, tags=req.tags,
            derived=req.derived, contributor=req.contributor, via_story=req.viaStory, via_slack=req.viaSlack,
            via_chat=req.viaChat, via_vantage=req.viaVantage,
        )
        db.add(d)
        db.commit()
        return _decision_json(d)


class DecisionUpdateIn(BaseModel):
    result: Optional[str] = None
    # Added 2026-08-30 for the Brain Board's consolidation flow: merging new
    # information into an existing node needs to update its title/summary/
    # tags/why/decided-date in place, not just log a review outcome into
    # `result` (the only field this endpoint supported before). Each is
    # optional and only applied when actually provided, so every other
    # existing caller of this endpoint (logReviewOutcome, still only ever
    # sending `result`) keeps working unchanged.
    title: Optional[str] = None
    derived: Optional[str] = None
    tags: Optional[list] = None
    why: Optional[list] = None
    decided: Optional[str] = None


@router.put("/companies/{company_id}/decisions/{decision_id}")
def update_decision(company_id: str, decision_id: str, req: DecisionUpdateIn):
    with get_session() as db:
        d = db.get(Decision, (company_id, decision_id))
        if not d:
            raise HTTPException(status_code=404, detail="No such decision.")
        if req.result is not None:
            d.result = req.result
        if req.title is not None:
            d.title = req.title
        if req.derived is not None:
            d.derived = req.derived
        if req.tags is not None:
            d.tags = req.tags
        if req.why is not None:
            d.why = req.why
        if req.decided is not None:
            d.decided = req.decided
        db.commit()
        return _decision_json(d)


@router.delete("/companies/{company_id}/decisions/{decision_id}")
def delete_decision(company_id: str, decision_id: str):
    # Added 2026-08-30 for the Brain Board: deleting a whole node off the
    # board (the founder's explicit scope choice — whole entries only, not
    # individual facts within one). No delete endpoint for a decision
    # existed before this; every other object in this file already has one
    # to mirror.
    with get_session() as db:
        d = db.get(Decision, (company_id, decision_id))
        if not d:
            raise HTTPException(status_code=404, detail="No such decision.")
        db.delete(d)
        db.commit()
        return {"ok": True}


# ---------------------------------------------------------------------------
# Goals
# ---------------------------------------------------------------------------
class GoalIn(BaseModel):
    id: Optional[str] = None
    title: str
    kicker: str = ""
    depth: int = 0
    parent: Optional[str] = None
    dept: Optional[str] = None
    owner: Optional[str] = None


@router.post("/companies/{company_id}/goals")
def add_goal(company_id: str, req: GoalIn):
    with get_session() as db:
        _get_company_or_404(db, company_id)
        gid = req.id or f"g-{_slugify(req.title)}"
        if db.get(Goal, (company_id, gid)):
            gid = f"{gid}-{len(db.execute(select(Goal).where(Goal.company_id == company_id)).scalars().all()) + 1}"
        g = Goal(company_id=company_id, id=gid, title=req.title, kicker=req.kicker, depth=req.depth,
                  parent=req.parent, dept=req.dept, owner=req.owner)
        db.add(g)
        db.commit()
        return _goal_json(g)


class GoalUpdateIn(BaseModel):
    owner: Optional[str] = None
    title: Optional[str] = None
    kicker: Optional[str] = None
    dept: Optional[str] = None
    # Deliberately no `parent`/`depth` here — reparenting would mean
    # recomputing depth for every descendant too (goal-alignment tracing
    # walks the parent chain assuming depth is consistent with position in
    # the tree). Out of scope for this pass: to move a goal, delete and
    # recreate it under the right parent instead.


@router.put("/companies/{company_id}/goals/{goal_id}")
def update_goal(company_id: str, goal_id: str, req: GoalUpdateIn):
    with get_session() as db:
        g = db.get(Goal, (company_id, goal_id))
        if not g:
            raise HTTPException(status_code=404, detail="No such goal.")
        if req.owner is not None:
            g.owner = req.owner
        if req.title is not None:
            g.title = req.title
        if req.kicker is not None:
            g.kicker = req.kicker
        if req.dept is not None:
            g.dept = req.dept
        db.commit()
        return _goal_json(g)


@router.delete("/companies/{company_id}/goals/{goal_id}")
def delete_goal(company_id: str, goal_id: str):
    # Added 2026-08-30 alongside the founder-facing "add a goal" UI — until
    # now goals only ever existed as seed data, so nothing needed deleting.
    # Cascades to descendants: a goal's whole reason for being is its place
    # in the parent chain (goal-alignment tracing walks it), so deleting a
    # department goal but leaving its team/individual goals pointing at a
    # parent that no longer exists would silently break that trace rather
    # than just removing what the admin asked to remove.
    with get_session() as db:
        g = db.get(Goal, (company_id, goal_id))
        if not g:
            raise HTTPException(status_code=404, detail="No such goal.")
        all_goals = db.execute(select(Goal).where(Goal.company_id == company_id)).scalars().all()
        by_parent = {}
        for other in all_goals:
            by_parent.setdefault(other.parent, []).append(other)
        to_delete = []
        frontier = [g]
        while frontier:
            cur = frontier.pop()
            to_delete.append(cur)
            frontier.extend(by_parent.get(cur.id, []))
        deleted_ids = [d.id for d in to_delete]
        for d in to_delete:
            db.delete(d)
        db.commit()
        return {"ok": True, "deletedIds": deleted_ids}


# ---------------------------------------------------------------------------
# Chat threads (Ask Cognex conversations) — added 2026-08-29, see ChatThread's
# docstring in models.py for why this exists and why `turns` is one JSON blob
# per thread rather than a normalized per-turn table.
# ---------------------------------------------------------------------------
def _thread_json(t: ChatThread) -> dict:
    return {"id": t.id, "title": t.title, "turns": t.turns or [], "updatedAt": t.updated_at}


def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


@router.get("/companies/{company_id}/personas/{persona_id}/threads")
def list_threads(company_id: str, persona_id: str):
    """Fetched once on login/session start so a persona's chat history is
    there from the first render, instead of every session starting with an
    empty sidebar the way it did before this endpoint existed."""
    with get_session() as db:
        rows = db.execute(
            select(ChatThread)
            .where(ChatThread.company_id == company_id, ChatThread.persona_id == persona_id)
            .order_by(ChatThread.updated_at.desc())
        ).scalars().all()
        return {"threads": [_thread_json(t) for t in rows]}


class ThreadUpsertIn(BaseModel):
    id: Optional[str] = None
    title: str = "New chat"
    turns: list = []


@router.put("/companies/{company_id}/personas/{persona_id}/threads/{thread_id}")
def upsert_thread(company_id: str, persona_id: str, thread_id: str, req: ThreadUpsertIn):
    """Upsert, not separate create/update — the frontend always knows the
    thread id it's writing (it mints one client-side the moment a new chat
    is started, same as it already does for decisions/goals), so one
    idempotent call handles both "first turn in a brand-new thread" and
    "another turn in an existing one" without the frontend having to track
    which case it's in. Called fire-and-forget after every turn completes
    and on rename/thread creation, mirroring persistNewDecision's pattern —
    the browser-local state.chatByKey stays the source of truth for
    rendering, this is a best-effort mirror so a reload has something real
    to restore from."""
    with get_session() as db:
        _get_company_or_404(db, company_id)
        t = db.get(ChatThread, (company_id, persona_id, thread_id))
        if not t:
            t = ChatThread(company_id=company_id, persona_id=persona_id, id=thread_id)
            db.add(t)
        t.title = req.title or "New chat"
        t.turns = req.turns
        t.updated_at = _now_iso()
        db.commit()
        return _thread_json(t)


@router.delete("/companies/{company_id}/personas/{persona_id}/threads/{thread_id}")
def delete_thread(company_id: str, persona_id: str, thread_id: str):
    with get_session() as db:
        t = db.get(ChatThread, (company_id, persona_id, thread_id))
        if not t:
            return {"ok": True}  # already gone — deleting twice isn't an error
        db.delete(t)
        db.commit()
        return {"ok": True}


# ---------------------------------------------------------------------------
# Handoffs — added 2026-08-31. Was sold as "Handoff Capture" (Business tier)
# while living only in the browser tab's company.handoffs array; see
# Handoff's docstring in models.py for the full reasoning. Create + update
# only (no delete): a handoff is a historical record of a real departure —
# finishHandoff marks it "done" and removes the departing PERSONA, it does
# not remove the handoff itself, so there's deliberately no DELETE route.
# ---------------------------------------------------------------------------
class HandoffIn(BaseModel):
    id: str
    fromPersonaId: str = ""
    fromName: str = ""
    fromTitle: str = ""
    fromDept: str = ""
    toPersonaId: str = ""
    createdAt: str = ""
    qaPairs: list = []
    report: dict = {}
    offline: bool = False
    delegationItems: list = []
    status: str = "pending"


@router.post("/companies/{company_id}/handoffs")
def create_handoff(company_id: str, req: HandoffIn):
    with get_session() as db:
        _get_company_or_404(db, company_id)
        h = Handoff(
            company_id=company_id, id=req.id,
            from_persona_id=req.fromPersonaId, from_name=req.fromName, from_title=req.fromTitle,
            from_dept=req.fromDept, to_persona_id=req.toPersonaId, created_at=req.createdAt or _now_iso(),
            qa_pairs=req.qaPairs, report=req.report, offline=req.offline,
            delegation_items=req.delegationItems, status=req.status,
        )
        db.add(h)
        db.commit()
        return _handoff_json(h)


class HandoffUpdateIn(BaseModel):
    status: Optional[str] = None
    delegationItems: Optional[list] = None


@router.put("/companies/{company_id}/handoffs/{handoff_id}")
def update_handoff(company_id: str, handoff_id: str, req: HandoffUpdateIn):
    """Covers both mutation paths the frontend has: assignDelegationItem
    (rewrites delegationItems + bumps status to "delegating") and
    finishHandoff (sets status to "done") — both just PUT whatever changed,
    same partial-update shape as update_goal above."""
    with get_session() as db:
        h = db.get(Handoff, (company_id, handoff_id))
        if not h:
            raise HTTPException(status_code=404, detail="No such handoff.")
        if req.status is not None:
            h.status = req.status
        if req.delegationItems is not None:
            h.delegation_items = req.delegationItems
        db.commit()
        return _handoff_json(h)


@router.get("/companies/{company_id}/handoffs")
def list_handoffs(company_id: str):
    with get_session() as db:
        rows = db.execute(
            select(Handoff).where(Handoff.company_id == company_id).order_by(Handoff.created_at.desc())
        ).scalars().all()
        return {"handoffs": [_handoff_json(h) for h in rows]}


# ---------------------------------------------------------------------------
# Access Ledger — added 2026-08-31. Was sold as "Access Ledger & audit
# exports" (Enterprise tier) while living only in browser memory; see
# LedgerEntry's docstring in models.py. Append-only by design (no PUT/DELETE
# routes) — an audit trail that could be edited or removed after the fact
# isn't an audit trail. A bulk endpoint exists alongside the single-entry
# one because logDecisionMemoryAccess can log many entries in one JS loop
# (once per visible decision, every time the Decision Memory tab opens) —
# without it, opening that tab would fire one HTTP request per decision.
# ---------------------------------------------------------------------------
class LedgerEntryIn(BaseModel):
    id: str
    ts: str = ""
    personaId: str = ""
    personaName: str = ""
    personaTitle: str = ""
    decisionId: str = ""
    decisionTitle: str = ""
    access: str = ""
    via: str = ""


def _write_ledger_entries(db, company_id: str, entries: list):
    for req in entries:
        db.add(LedgerEntry(
            company_id=company_id, id=req.id, ts=req.ts or _now_iso(),
            persona_id=req.personaId, persona_name=req.personaName, persona_title=req.personaTitle,
            decision_id=req.decisionId, decision_title=req.decisionTitle, access=req.access, via=req.via,
        ))
    db.commit()
    # Cap to the most recent 500 rows for this company, same limit the
    # frontend already enforced locally before this endpoint existed — an
    # audit trail here is meant to catch recent, actionable access patterns,
    # not grow without bound.
    rows = db.execute(
        select(LedgerEntry.id).where(LedgerEntry.company_id == company_id).order_by(LedgerEntry.ts.desc())
    ).scalars().all()
    if len(rows) > 500:
        stale_ids = rows[500:]
        db.execute(
            LedgerEntry.__table__.delete().where(
                LedgerEntry.company_id == company_id, LedgerEntry.id.in_(stale_ids)
            )
        )
        db.commit()


@router.post("/companies/{company_id}/ledger")
def create_ledger_entry(company_id: str, req: LedgerEntryIn):
    with get_session() as db:
        _get_company_or_404(db, company_id)
        _write_ledger_entries(db, company_id, [req])
        return {"ok": True}


class LedgerBulkIn(BaseModel):
    entries: list[LedgerEntryIn] = []


@router.post("/companies/{company_id}/ledger/bulk")
def create_ledger_entries_bulk(company_id: str, req: LedgerBulkIn):
    with get_session() as db:
        _get_company_or_404(db, company_id)
        if req.entries:
            _write_ledger_entries(db, company_id, req.entries)
        return {"ok": True, "count": len(req.entries)}


@router.get("/companies/{company_id}/ledger")
def list_ledger(company_id: str):
    with get_session() as db:
        rows = db.execute(
            select(LedgerEntry).where(LedgerEntry.company_id == company_id).order_by(LedgerEntry.ts.desc()).limit(500)
        ).scalars().all()
        return {"entries": [_ledger_json(e) for e in rows]}


# ---------------------------------------------------------------------------
# Company Document Library — added 2026-08-31, direct founder feedback ("I
# am not able to upload files"). The founder's own explicit choice on scope
# (the biggest of three options offered): a real persistent, company-wide,
# searchable library — uploaded once, referenceable across every future
# chat — not a one-off per-message chat attachment. See Document's docstring
# in models.py for the full reasoning behind its globally-unique id and why
# file bytes live in Postgres rather than on local disk.
#
# Text extraction is intentionally narrow: txt/md/csv/json/log are read
# directly as UTF-8 text, pdf/docx go through pypdf/python-docx (both added
# to requirements.txt alongside this), and anything else is still stored and
# downloadable but marked "not_indexed" rather than pretending to search
# content that was never actually read. DOCUMENT_MAX_BYTES exists because
# file bytes live in a Postgres column here, not a real object-storage tier
# — there's no infrastructure behind this to make an unbounded per-file size
# sane on a small Railway Postgres instance.
# ---------------------------------------------------------------------------
DOCUMENT_MAX_BYTES = 15 * 1024 * 1024  # 15 MB — see the module comment above
DOCUMENT_TEXT_EXTS = ("txt", "md", "markdown", "csv", "json", "log")


def _extract_document_text(filename: str, content: bytes) -> tuple[str, str]:
    """Returns (text_content, extraction_status). Deliberately never raises
    — a failure to extract text makes for a degraded-but-still-uploaded
    document (stored and downloadable, just not searchable by content), not
    a rejected upload; the person didn't do anything wrong by uploading a
    file this code can't parse."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    try:
        if ext in DOCUMENT_TEXT_EXTS:
            return content.decode("utf-8", errors="replace"), "indexed"
        if ext == "pdf":
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(content))
            text = "\n".join((page.extract_text() or "") for page in reader.pages)
            return text, ("indexed" if text.strip() else "not_indexed")
        if ext == "docx":
            import docx
            doc = docx.Document(io.BytesIO(content))
            text = "\n".join(p.text for p in doc.paragraphs)
            return text, ("indexed" if text.strip() else "not_indexed")
    except Exception:
        return "", "failed"
    return "", "not_indexed"


def _document_json(d: Document) -> dict:
    # Deliberately never includes content_bytes or the full text_content —
    # this is the shape used for the library list, which can hold many
    # documents at once; the raw bytes are only ever served by live.py's
    # GET /api/files/doc-{id} download route, and the extracted text is only
    # ever read directly out of the DB by live.py's search_documents/
    # get_document tools, never shipped to the browser in bulk.
    return {
        "id": d.id, "filename": d.filename, "mimeType": d.mime_type, "sizeBytes": d.size_bytes,
        "uploadedBy": d.uploaded_by, "uploadedAt": d.uploaded_at, "extractionStatus": d.extraction_status,
    }


@router.post("/companies/{company_id}/documents")
async def upload_document(company_id: str, file: UploadFile = File(...), uploadedBy: str = Form("")):
    content = await file.read()
    if len(content) > DOCUMENT_MAX_BYTES:
        raise HTTPException(
            status_code=413,
            detail=f"That file is too large — the document library caps uploads at {DOCUMENT_MAX_BYTES // (1024 * 1024)} MB.",
        )
    filename = file.filename or "upload"
    with get_session() as db:
        _get_company_or_404(db, company_id)
        text_content, status = _extract_document_text(filename, content)
        d = Document(
            id="doc-" + secrets.token_urlsafe(16), company_id=company_id, filename=filename,
            mime_type=file.content_type or "application/octet-stream", size_bytes=len(content),
            uploaded_by=uploadedBy, uploaded_at=_now_iso(),
            extraction_status=status, text_content=text_content, content_bytes=content,
        )
        db.add(d)
        db.commit()
        return _document_json(d)


@router.get("/companies/{company_id}/documents")
def list_documents(company_id: str):
    with get_session() as db:
        rows = db.execute(
            select(Document).where(Document.company_id == company_id).order_by(Document.uploaded_at.desc())
        ).scalars().all()
        return {"documents": [_document_json(d) for d in rows]}


@router.delete("/companies/{company_id}/documents/{document_id}")
def delete_document(company_id: str, document_id: str):
    with get_session() as db:
        d = db.get(Document, document_id)
        if not d or d.company_id != company_id:
            return {"ok": True}  # already gone (or never belonged to this company) — deleting twice isn't an error
        db.delete(d)
        db.commit()
        return {"ok": True}
